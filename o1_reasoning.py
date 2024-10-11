import streamlit as st
import openai
from PIL import Image
import requests
from io import BytesIO
from docx import Document
from fpdf import FPDF

# Set up the AIMLAPI client for text and image generation
openai.api_key = "Your AIMLAPI Key"
openai.api_base = "https://api.aimlapi.com"

# Mapping of user-friendly names to AIMLAPI model identifiers
model_map = {
    "dall-e": "dall-e-3",  # Replace with "dall-e-2" if needed
    "stable-diffusion": "stabilityai/stable-diffusion-2-1",
    "flux": "flux-pro"
}

# Define function to summarize research paper
def summarize_paper(paper_text):
    response = openai.ChatCompletion.create(
        model="o1-mini",
        messages=[
            {"role": "user", "content": f"Summarize the following research paper: {paper_text}"}
        ],
        max_tokens=1000,
    )
    summary = response.choices[0].message['content']
    return summary

# Define function to generate sections of a research paper
def write_research_paper(title, existing_paper_summary):
    prompts = {
        "abstract": f"Write an abstract for a research paper titled '{title}' based on the following summary: {existing_paper_summary}",
        "introduction": f"Write an introduction for a research paper titled '{title}' based on the following summary: {existing_paper_summary}",
        "methodology": f"Describe the methodology for a research paper titled '{title}' based on the following summary: {existing_paper_summary}",
        "results": f"Discuss the results for a research paper titled '{title}' based on the following summary: {existing_paper_summary}",
        "conclusion": f"Conclude a research paper titled '{title}' based on the following summary: {existing_paper_summary}",
        "use_cases": f"Provide the use cases related to the research paper titled '{title}' based on the following summary: {existing_paper_summary}"
    }
    
    paper_sections = {}
    for section, prompt in prompts.items():
        response = openai.ChatCompletion.create(
            model="o1-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
        )
        paper_sections[section] = response.choices[0].message['content']
    
    return paper_sections

# Function to generate process flow diagrams based on the methodology
def generate_process_flow(content):
    flow_prompt = f"Create a process flow diagram based on the following methodology: {content}"
    response = openai.Image.create(
        model="dall-e-3",
        prompt=flow_prompt,
        n=1,
        size="1024x1024"
    )
    try:
        image_url = response['data'][0]['url']
        return image_url
    except KeyError:
        st.error("Error generating flowchart. Please try again.")
        return None

# Function to create a DOCX file
def create_docx(paper_sections, title):
    doc = Document()
    doc.add_heading(title, 0)
    for section, content in paper_sections.items():
        doc.add_heading(section.capitalize(), level=1)
        doc.add_paragraph(content)
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# Function to create a simple PDF
def create_pdf(paper_sections, title):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    pdf.cell(200, 10, title, ln=True, align='C')
    for section, content in paper_sections.items():
        pdf.set_font("Arial", size=14)
        pdf.cell(200, 10, section.capitalize(), ln=True)
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, content)
    
    pdf_stream = BytesIO()
    pdf.output(pdf_stream)
    pdf_stream.seek(0)
    return pdf_stream

# Streamlit App Layout
st.title("Research Paper Generation Tool with Visualizations")

# Input for research paper content
paper_text = st.text_area("Enter the content of the research paper:", height=300)
title = st.text_input("Enter the title for the new research paper:")

# Button to generate summary and research paper
if st.button("Generate Research Paper"):
    if paper_text and title:
        summary = summarize_paper(paper_text)
        new_paper = write_research_paper(title, summary)
        
        # Display generated research paper in a box
        with st.expander("Generated Research Paper", expanded=True):
            for section, content in new_paper.items():
                st.write(f"**{section.capitalize()}:**")
                st.markdown(content)  # Use markdown for formatting
            
                # Generate and display process flow diagrams for methodology
                if section == "methodology":
                    flow_url = generate_process_flow(content)
                    if flow_url:
                        image = Image.open(requests.get(flow_url, stream=True).raw)
                        st.image(image, caption="Methodology Process Flow Diagram", use_column_width=True)
                
                # Generate and display use case diagrams
                if section == "use_cases":
                    flow_url = generate_process_flow(content)
                    if flow_url:
                        image = Image.open(requests.get(flow_url, stream=True).raw)
                        st.image(image, caption="Use Case Diagram", use_column_width=True)

        # Download options
        if st.button("Download as DOCX"):
            doc_stream = create_docx(new_paper, title)
            st.download_button("Download Research Paper (DOCX)", doc_stream, file_name=f"{title}.docx")

        if st.button("Download as PDF"):
            pdf_stream = create_pdf(new_paper, title)
            st.download_button("Download Research Paper (PDF)", pdf_stream, file_name=f"{title}.pdf")
    else:
        st.error("Please provide both the research paper content and a title for the new paper.")
